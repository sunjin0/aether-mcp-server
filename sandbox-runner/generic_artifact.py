"""平台维护的常用办公产物离线渲染器。

Agent 只能提供内容，不能提供可执行代码。该脚本冻结在 runner 镜像中，是通用产物
任务唯一可使用的入口。
"""
import html
import json
import os
import re
from html.parser import HTMLParser
from pathlib import Path


def safe_name(value: str, extension: str) -> str:
    base = re.sub(r"[\\/:*?\"<>|]+", "-", value or "generated").strip(" .-")
    base = base[:100] or "generated"
    return base if base.lower().endswith("." + extension) else base + "." + extension


def inferred_pdf_name(title: str, content: str) -> str:
    """Use the model title as the base name and enrich resume titles when possible."""
    plain = html.unescape(re.sub(r"<[^>]+>", " ", content or ""))
    intent = re.search(r"求职意向\s*[：:]\s*([^\n<]+)", plain)
    if not intent or "简历" in (title or ""):
        return title or "generated"
    role = re.split(r"[／/｜|]", intent.group(1), maxsplit=1)[0]
    role = re.sub(r"[（(].*?[）)]", "", role).strip(" ：:-")
    return f"{title}-简历-{role}" if role else title


def markdown_lines(content: str):
    for line in content.replace("\r\n", "\n").split("\n"):
        yield line.rstrip()


def is_duplicate_title_heading(line: str, title: str) -> bool:
    """仅当首行 Markdown H1 与产物标题重复时返回 True。"""
    heading = re.match(r"^#\s+(.+?)\s*$", line or "")
    return heading is not None and heading.group(1).strip() == (title or "").strip()


def strip_duplicate_html_title(content: str, title: str) -> str:
    """文档 H1 由渲染器生成，因此移除正文开头重复的 H1。"""
    escaped_title = re.escape(html.escape((title or "").strip()))
    if not escaped_title:
        return content
    return re.sub(r"^\s*<h1(?:\s[^>]*)?>\s*" + escaped_title + r"\s*</h1>\s*", "", content,
                  count=1, flags=re.IGNORECASE)


def parse_table(lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def inline_markdown(value: str) -> str:
    """渲染产物生成允许的安全行内 Markdown 子集。"""
    text = html.escape(value or "")
    # 先转义输入，再仅引入渲染器生成的标签，使原始 HTML 失效但保留 PDF 所需的
    # Markdown 格式。
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*|__([^_]+)__", lambda match: "<strong>" + (match.group(1) or match.group(2)) + "</strong>", text)
    text = re.sub(r"~~([^~]+)~~", r"<del>\1</del>", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)|(?<!_)_([^_\n]+)_(?!_)", lambda match: "<em>" + (match.group(1) or match.group(2)) + "</em>", text)

    def link(match):
        label, url = match.group(1), html.unescape(match.group(2)).strip()
        if re.match(r"https?://[^\s<]+$", url, re.IGNORECASE):
            return '<a href="' + html.escape(url, quote=True) + '">' + label + "</a>"
        return label

    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", link, text)


HTML_TAGS = {"html", "head", "body", "main", "header", "footer", "section", "article", "aside", "h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "span", "small", "strong", "b", "em", "i", "u", "s", "del", "br", "ul", "ol", "li", "table", "thead", "tbody", "tr", "th", "td", "blockquote", "hr", "pre", "code", "a", "style"}
VOID_HTML_TAGS = {"br", "hr"}
BLOCKED_HTML_TAGS = {"script", "iframe", "object", "embed", "svg", "math", "form", "input", "button", "video", "audio", "canvas", "title", "img", "link"}
BLOCKED_VOID_HTML_TAGS = {"img", "link", "input", "embed"}
SAFE_STYLE_PROPERTIES = {"color", "background", "background-color", "font-family", "font-size", "font-weight", "font-style", "font-variant", "text-align", "text-indent", "text-decoration", "text-transform", "letter-spacing", "word-spacing", "word-break", "overflow", "overflow-x", "overflow-y", "overflow-wrap", "box-sizing", "margin", "margin-top", "margin-right", "margin-bottom", "margin-left", "padding", "padding-top", "padding-right", "padding-bottom", "padding-left", "border", "border-top", "border-right", "border-bottom", "border-left", "border-radius", "border-collapse", "border-spacing", "box-shadow", "width", "min-width", "max-width", "height", "min-height", "max-height", "line-height", "white-space", "vertical-align", "display", "position", "z-index", "top", "right", "bottom", "left", "content", "flex", "flex-direction", "flex-wrap", "flex-grow", "flex-shrink", "flex-basis", "justify-content", "align-items", "align-content", "gap", "row-gap", "column-gap", "grid-template-columns", "grid-template-rows", "grid-column", "grid-row", "list-style", "list-style-type", "page-break-before", "page-break-after", "page-break-inside", "break-before", "break-after", "break-inside", "opacity", "size", "print-color-adjust", "-webkit-print-color-adjust"}


def sanitize_style(value: str) -> str:
    safe = []
    for declaration in (value or "").split(";"):
        if ":" not in declaration:
            continue
        name, css_value = (part.strip() for part in declaration.split(":", 1))
        lowered = css_value.lower()
        if name.lower() not in SAFE_STYLE_PROPERTIES or any(token in lowered for token in ("url(", "@import", "expression(", "javascript:", "behavior:")):
            continue
        safe.append(name.lower() + ":" + css_value)
    return ";".join(safe)


def sanitize_css(value: str) -> str:
    """Keep safe model CSS, including print-only layout rules, without external loading."""
    value = re.sub(r"/\*.*?\*/", "", value or "", flags=re.DOTALL)
    rules, cursor = [], 0
    while cursor < len(value):
        opening = value.find("{", cursor)
        if opening < 0:
            break
        selector = value[cursor:opening].strip()
        depth, closing = 1, opening + 1
        while closing < len(value) and depth:
            if value[closing] == "{":
                depth += 1
            elif value[closing] == "}":
                depth -= 1
            closing += 1
        if depth:
            break
        declarations = value[opening + 1:closing - 1]
        selector_lower = selector.lower()
        if any(token in selector_lower for token in ("url(", "expression", "javascript:")):
            cursor = closing
            continue
        if selector_lower == "@page":
            cleaned = sanitize_style(declarations)
        elif re.fullmatch(r"@media\s+print", selector_lower):
            cleaned = sanitize_css(declarations)
        elif selector.startswith("@"):
            cleaned = ""
        else:
            cleaned = sanitize_style(declarations)
        if cleaned:
            rules.append(selector + "{" + cleaned + "}")
        cursor = closing
    return "".join(rules)


class SafeHtml(HTMLParser):
    """仅保留布局 HTML，丢弃代码、外部资源和主动属性。"""
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts, self.blocked_depth, self.style_depth = [], 0, 0

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in BLOCKED_HTML_TAGS:
            if tag not in BLOCKED_VOID_HTML_TAGS:
                self.blocked_depth += 1
            return
        if self.blocked_depth or tag not in HTML_TAGS:
            return
        allowed = []
        for name, value in attrs:
            name, value = name.lower(), value or ""
            if name == "style":
                cleaned = sanitize_style(value)
                if cleaned:
                    allowed.append((name, cleaned))
            elif name == "class":
                allowed.append((name, value))
            elif tag == "a" and name == "href" and re.fullmatch(r"https?://[^\s<>]+", value, re.IGNORECASE):
                allowed.append((name, value))
        rendered = "".join(" %s=\"%s\"" % (name, html.escape(value, quote=True)) for name, value in allowed)
        self.parts.append("<" + tag + rendered + ">")
        if tag == "style":
            self.style_depth += 1

    def handle_startendtag(self, tag, attrs):
        self.handle_starttag(tag, attrs)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in BLOCKED_HTML_TAGS:
            self.blocked_depth = max(0, self.blocked_depth - 1)
            return
        if tag == "style":
            self.style_depth = max(0, self.style_depth - 1)
        if not self.blocked_depth and tag in HTML_TAGS and tag not in VOID_HTML_TAGS:
            self.parts.append("</" + tag + ">")

    def handle_data(self, data):
        if not self.blocked_depth:
            self.parts.append(sanitize_css(data) if self.style_depth else html.escape(data))


def looks_like_html(content: str) -> bool:
    return bool(re.search(r"<!doctype\s+html|<(?:html|body|h[1-6]|p|div|span|table|ul|ol|blockquote)\b", content or "", re.IGNORECASE))


def normalize_html_content(content: str) -> str:
    """Decode an HTML string that was accidentally JSON-escaped a second time."""
    value = content or ""
    if "\\n" not in value and '\\"' not in value and "\\t" not in value:
        return value
    return (value.replace("\\r\\n", "\n").replace("\\n", "\n")
            .replace("\\t", "\t").replace('\\"', '"'))


def sanitize_html(content: str) -> str:
    parser = SafeHtml()
    parser.feed(normalize_html_content(content))
    parser.close()
    return "".join(parser.parts)


def render_docx(title: str, content: str, output: Path):
    from docx import Document
    document = Document()
    document.add_heading(title, 0)
    lines = list(markdown_lines(content))
    index = 0
    while index < len(lines):
        line = lines[index]
        if index == 0 and is_duplicate_title_heading(line, title):
            index += 1
            continue
        if line.startswith("#"):
            heading = re.match(r"^(#{1,6})\s*(.*)$", line)
            if heading:
                document.add_heading(heading.group(2), min(len(heading.group(1)), 4))
            index += 1
            continue
        if line.strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index]); index += 1
            rows = parse_table(table_lines)
            if rows:
                table = document.add_table(rows=0, cols=max(len(row) for row in rows))
                table.style = "Table Grid"
                for row in rows:
                    cells = table.add_row().cells
                    for col, value in enumerate(row): cells[col].text = value
            continue
        if line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line.strip():
            document.add_paragraph(line)
        index += 1
    document.save(output)


def render_xlsx(title: str, content: str, document_plan: dict, output: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Font
    workbook = Workbook(); workbook.remove(workbook.active)
    sheets = document_plan.get("sheets") if isinstance(document_plan, dict) else None
    if not isinstance(sheets, list) or not sheets:
        rows = parse_table(list(markdown_lines(content))) or [[line] for line in markdown_lines(content) if line.strip()]
        sheets = [{"name": title[:31] or "Sheet1", "rows": rows}]
    for sheet_def in sheets:
        worksheet = workbook.create_sheet(str(sheet_def.get("name") or "Sheet")[:31])
        rows = sheet_def.get("rows") or []
        for row in rows: worksheet.append(row if isinstance(row, list) else [row])
        if worksheet.max_row:
            for cell in worksheet[1]: cell.font = Font(bold=True)
        for column in worksheet.columns:
            width = min(48, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
            worksheet.column_dimensions[column[0].column_letter].width = width
    workbook.save(output)


def render_pdf(content: str, output: Path):
    from weasyprint import HTML
    if looks_like_html(content):
        render_pdf_html(content, output)
        return
    blocks, lines, index = [], list(markdown_lines(content)), 0
    unordered = re.compile(r"^\s*[-+*]\s+(.+)$")
    ordered = re.compile(r"^\s*(?:\d+[.)]|[一二三四五六七八九十]+、)\s*(.+)$")
    while index < len(lines):
        line = lines[index]
        if re.fullmatch(r"\s*(?:---+|\*\*\*+)\s*", line):
            blocks.append("<hr>"); index += 1; continue
        if line.strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index]); index += 1
            rows = parse_table(table_lines)
            if rows:
                header, body = rows[0], rows[1:]
                blocks.append("<table><thead><tr>" + "".join("<th>" + inline_markdown(cell) + "</th>" for cell in header) + "</tr></thead><tbody>" + "".join("<tr>" + "".join("<td>" + inline_markdown(cell) + "</td>" for cell in row) + "</tr>" for row in body) + "</tbody></table>")
            continue
        if unordered.match(line):
            items = []
            while index < len(lines) and unordered.match(lines[index]):
                items.append("<li>" + inline_markdown(unordered.match(lines[index]).group(1)) + "</li>"); index += 1
            blocks.append("<ul>" + "".join(items) + "</ul>"); continue
        if ordered.match(line):
            items = []
            while index < len(lines) and ordered.match(lines[index]):
                items.append("<li>" + inline_markdown(ordered.match(lines[index]).group(1)) + "</li>"); index += 1
            blocks.append("<ol>" + "".join(items) + "</ol>"); continue
        heading = re.match(r"^(#{1,4})\s*(.*)$", line)
        if heading:
            level = len(heading.group(1)); blocks.append(f"<h{level}>" + inline_markdown(heading.group(2)) + f"</h{level}>")
        elif line.strip(): blocks.append("<p>" + inline_markdown(line) + "</p>")
        index += 1
    style = """@page{size:A4;margin:18mm 15mm}body{font-family:'Noto Sans CJK SC','Noto Sans',sans-serif;color:#1f2937;font-size:10.5pt;line-height:1.65}h1{font-size:22pt;margin:0 0 16pt;border-bottom:2px solid #2563eb;padding-bottom:8pt}h2{font-size:16pt;margin:20pt 0 8pt;color:#1e3a8a}h3{font-size:13pt;margin:15pt 0 6pt}p{margin:0 0 7pt}ul,ol{margin:4pt 0 9pt;padding-left:20pt}hr{border:0;border-top:1px solid #cbd5e1;margin:12pt 0}strong{font-weight:700}em{font-style:italic}del{color:#6b7280;text-decoration:line-through}code{font-family:monospace;background:#f1f5f9;border-radius:3pt;padding:1pt 3pt;color:#9f1239}a{color:#2563eb;text-decoration:underline}table{width:100%;border-collapse:collapse;margin:10pt 0 14pt;font-size:8.5pt;table-layout:auto;page-break-inside:auto}thead{display:table-header-group;background:#eaf2ff}tr{page-break-inside:avoid}th,td{border:1px solid #9ca3af;padding:5pt 6pt;vertical-align:top;word-break:break-word}th{font-weight:700;color:#173b75}"""
    HTML(string="<meta charset='utf-8'><style>" + style + "</style>" + "".join(blocks)).write_pdf(output)


def render_pdf_html(content: str, output: Path):
    """移除活动内容和可远程加载部分后，才渲染 Agent 提供的 HTML。"""
    from weasyprint import HTML
    style = """@page{size:A4;margin:14mm}html,body{margin:0;padding:0}body{font-family:'Noto Sans CJK SC','Noto Sans',sans-serif;color:#1f2937;font-size:10.5pt;line-height:1.5}*,*:before,*:after{box-sizing:border-box}table{max-width:100%;border-collapse:collapse}img{max-width:100%}pre{white-space:pre-wrap;overflow-wrap:anywhere}td,th{overflow-wrap:anywhere}"""
    safe_body = sanitize_html(content)
    HTML(string="<meta charset='utf-8'><style>" + style + "</style>" + safe_body).write_pdf(output)


def main():
    payload = json.loads(os.environ.get("AETHER_INPUT_JSON") or "{}")
    format_name = str(payload.get("format") or "").lower()
    if format_name not in {"docx", "xlsx", "pdf"}: raise ValueError("unsupported artifact format")
    title = str(payload.get("title") or "generated")
    output = Path(os.environ["AETHER_OUTPUT_DIR"]); output.mkdir(parents=True, exist_ok=True)
    requested_name = str(payload.get("fileName") or "")
    if not requested_name and format_name == "pdf":
        requested_name = inferred_pdf_name(title, str(payload.get("content") or ""))
    target = output / safe_name(requested_name or title, format_name)
    content = str(payload.get("content") or "")
    if format_name == "docx": render_docx(title, content, target)
    elif format_name == "xlsx": render_xlsx(title, content, payload.get("document") or {}, target)
    else: render_pdf(content, target)
    print("generated=" + target.name)


if __name__ == "__main__": main()
